"""Block-aware document parsing for governed RAG ingestion.

The parser preserves tables as atomic blocks. PDF OCR is optional so the
offline profile stays deterministic; when optional dependencies are installed,
low-text PDF pages use a rasterize -> OCR fallback.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from pypdf import PdfReader


@dataclass(frozen=True, slots=True)
class DocumentBlock:
    block_type: str
    text: str
    page: int | None = None
    table_id: str | None = None
    row_start: int | None = None
    row_end: int | None = None
    source: str = "native"


@dataclass(frozen=True, slots=True)
class ParsedDocument:
    filename: str
    blocks: tuple[DocumentBlock, ...]
    parser: str
    warnings: tuple[str, ...] = ()

    @property
    def text(self) -> str:
        return "\n\n".join(block_to_markdown(block) for block in self.blocks)

    def diagnostics(self) -> dict[str, Any]:
        return {
            "parser": self.parser,
            "block_count": len(self.blocks),
            "table_count": sum(block.block_type == "table" for block in self.blocks),
            "ocr_block_count": sum(block.source == "ocr" for block in self.blocks),
            "warnings": list(self.warnings),
        }


def block_to_markdown(block: DocumentBlock) -> str:
    if block.block_type == "table":
        table_id = block.table_id or "table-unknown"
        return (
            f"<!-- table:start id={table_id} page={block.page or ''} -->\n"
            f"{block.text.strip()}\n"
            "<!-- table:end -->"
        )
    prefix = f"[page:{block.page}] " if block.page else ""
    return prefix + block.text.strip()


class DocumentParser:
    def __init__(self, *, ocr_enabled: bool = True, min_native_chars: int = 40) -> None:
        self.ocr_enabled = ocr_enabled
        self.min_native_chars = min_native_chars

    def parse(self, filename: str, raw: bytes | str) -> ParsedDocument:
        if isinstance(raw, str):
            raw = raw.encode("utf-8")
        suffix = Path(filename).suffix.casefold()
        if suffix == ".pdf":
            return self._parse_pdf(filename, raw)
        if suffix in {".docx"}:
            return self._parse_docx(filename, raw)
        if suffix in {".xlsx", ".xlsm"}:
            return self._parse_xlsx(filename, raw)
        if suffix in {".html", ".htm"}:
            return self._parse_html(filename, raw)
        if suffix in {".md", ".txt", ".csv"}:
            return self._parse_markdown(filename, raw)
        if suffix in {".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff"}:
            return self._parse_image(filename, raw)
        return ParsedDocument(
            filename=filename,
            blocks=(DocumentBlock("paragraph", raw.decode("utf-8-sig", errors="replace")),),
            parser="plain_text",
        )

    def _parse_markdown(self, filename: str, raw: bytes) -> ParsedDocument:
        text = raw.decode("utf-8-sig", errors="replace")
        blocks: list[DocumentBlock] = []
        paragraphs: list[str] = []
        table_lines: list[str] = []
        table_index = 0

        def flush_paragraph() -> None:
            if paragraphs:
                content = "\n".join(paragraphs).strip()
                if content:
                    blocks.append(DocumentBlock("paragraph", content))
                paragraphs.clear()

        def flush_table() -> None:
            nonlocal table_index
            if table_lines:
                table_index += 1
                blocks.append(
                    DocumentBlock(
                        "table",
                        "\n".join(table_lines),
                        table_id=f"table-{table_index}",
                        row_start=1,
                        row_end=len(table_lines),
                    )
                )
                table_lines.clear()

        for raw_line in text.splitlines():
            line = raw_line.strip()
            if line.startswith("|") or (table_lines and "|" in line):
                flush_paragraph()
                table_lines.append(line)
            elif not line and table_lines:
                flush_table()
            else:
                if table_lines:
                    flush_table()
                paragraphs.append(raw_line)
        flush_table()
        flush_paragraph()
        return ParsedDocument(filename, tuple(blocks), "markdown_block_parser")

    def _parse_pdf(self, filename: str, raw: bytes) -> ParsedDocument:
        reader = PdfReader(io.BytesIO(raw))
        blocks: list[DocumentBlock] = []
        warnings: list[str] = []
        for page_number, page in enumerate(reader.pages, start=1):
            native = (page.extract_text() or "").strip()
            if len(native) >= self.min_native_chars:
                blocks.extend(_text_to_blocks(native, page=page_number, source="native"))
                continue
            ocr_blocks = self._ocr_pdf_page(raw, page_number)
            if ocr_blocks:
                blocks.extend(ocr_blocks)
            elif native:
                blocks.extend(_text_to_blocks(native, page=page_number, source="native"))
                warnings.append(f"page {page_number}: native text is sparse; OCR unavailable")
            else:
                warnings.append(f"page {page_number}: no native text and OCR unavailable")
        return ParsedDocument(filename, tuple(blocks), "pdf_native_ocr_fallback", tuple(warnings))

    def _ocr_pdf_page(self, raw: bytes, page_number: int) -> list[DocumentBlock]:
        if not self.ocr_enabled:
            return []
        try:
            import fitz  # type: ignore[import-not-found]
            import numpy as np  # type: ignore[import-not-found]
            from PIL import Image  # type: ignore[import-not-found]
            from rapidocr_onnxruntime import RapidOCR  # type: ignore[import-not-found]
        except ImportError:
            return []
        document = fitz.open(stream=raw, filetype="pdf")
        try:
            page = document.load_page(page_number - 1)
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image = Image.open(io.BytesIO(pixmap.tobytes("png"))).convert("RGB")
            result, _ = RapidOCR()(np.asarray(image))
        finally:
            document.close()
        lines = [
            str(item[1]).strip()
            for item in (result or [])
            if len(item) >= 3 and str(item[1]).strip()
        ]
        if not lines:
            return []
        text = "\n".join(lines)
        return _text_to_blocks(text, page=page_number, source="ocr")

    def _parse_docx(self, filename: str, raw: bytes) -> ParsedDocument:
        try:
            from docx import Document as DocxDocument  # type: ignore[import-not-found]
            from docx.table import Table  # type: ignore[import-not-found]
            from docx.text.paragraph import Paragraph  # type: ignore[import-not-found]
        except ImportError as exc:
            raise RuntimeError("python-docx is required for DOCX ingestion") from exc
        document = DocxDocument(io.BytesIO(raw))
        blocks: list[DocumentBlock] = []
        table_index = 0
        for child in document.element.body.iterchildren():
            if child.tag.endswith("}p"):
                text = Paragraph(child, document).text.strip()
                if text:
                    blocks.append(DocumentBlock("paragraph", text, source="native"))
            elif child.tag.endswith("}tbl"):
                table_index += 1
                table = Table(child, document)
                lines = []
                for row_number, row in enumerate(table.rows, start=1):
                    cells = [re.sub(r"\s+", " ", cell.text).strip() for cell in row.cells]
                    lines.append(f"| {row_number} | " + " | ".join(cells) + " |")
                if lines:
                    blocks.append(
                        DocumentBlock(
                            "table",
                            "| row | "
                            + " | ".join(
                                f"col{n}" for n in range(1, len(table.rows[0].cells) + 1)
                            )
                            + " |\n"
                            + "|---|"
                            + "---|" * len(table.rows[0].cells)
                            + "\n"
                            + "\n".join(lines),
                            table_id=f"table-{table_index}",
                            row_start=1,
                            row_end=len(table.rows),
                        )
                    )
        return ParsedDocument(filename, tuple(blocks), "python-docx")

    def _parse_xlsx(self, filename: str, raw: bytes) -> ParsedDocument:
        try:
            from openpyxl import load_workbook
            from openpyxl.utils import get_column_letter
        except ImportError as exc:
            raise RuntimeError("openpyxl is required for XLSX ingestion") from exc
        workbook = load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
        blocks: list[DocumentBlock] = []
        try:
            for sheet in workbook.worksheets:
                rows = []
                for row_number, row in enumerate(sheet.iter_rows(), start=1):
                    cells = [
                        f"{get_column_letter(index)}{row_number}={str(cell.value).strip()}"
                        for index, cell in enumerate(row, start=1)
                        if cell.value is not None
                    ]
                    if cells:
                        rows.append(f"Row {row_number}: " + " | ".join(cells))
                if rows:
                    blocks.append(
                        DocumentBlock(
                            "table",
                            f"Sheet: {sheet.title}\n" + "\n".join(rows),
                            table_id=f"sheet-{sheet.title}",
                            row_start=1,
                            row_end=len(rows),
                        )
                    )
        finally:
            workbook.close()
        return ParsedDocument(filename, tuple(blocks), "openpyxl")

    def _parse_html(self, filename: str, raw: bytes) -> ParsedDocument:
        text = raw.decode("utf-8-sig", errors="replace")
        tables = re.findall(r"<table[^>]*>(.*?)</table>", text, flags=re.IGNORECASE | re.DOTALL)
        blocks: list[DocumentBlock] = []
        for index, table in enumerate(tables, start=1):
            rows = re.findall(r"<tr[^>]*>(.*?)</tr>", table, flags=re.IGNORECASE | re.DOTALL)
            cells = []
            for row in rows:
                row_cells = re.findall(
                    r"<t[dh][^>]*>(.*?)</t[dh]>",
                    row,
                    flags=re.IGNORECASE | re.DOTALL,
                )
                cells.append(" | ".join(re.sub(r"<[^>]+>", "", cell).strip() for cell in row_cells))
            if cells:
                blocks.append(DocumentBlock("table", "\n".join(cells), table_id=f"table-{index}"))
        plain = re.sub(r"<[^>]+>", " ", text)
        plain = re.sub(r"\s+", " ", plain).strip()
        if plain:
            blocks.insert(0, DocumentBlock("paragraph", plain))
        return ParsedDocument(filename, tuple(blocks), "html_block_parser")

    def _parse_image(self, filename: str, raw: bytes) -> ParsedDocument:
        result = self._ocr_image(raw)
        if not result:
            return ParsedDocument(filename, (), "image_ocr", ("OCR unavailable or empty",))
        return ParsedDocument(filename, tuple(_text_to_blocks(result, source="ocr")), "image_ocr")

    def _ocr_image(self, raw: bytes) -> str:
        if not self.ocr_enabled:
            return ""
        try:
            import numpy as np  # type: ignore[import-not-found]
            from PIL import Image  # type: ignore[import-not-found]
            from rapidocr_onnxruntime import RapidOCR  # type: ignore[import-not-found]
        except ImportError:
            return ""
        with Image.open(io.BytesIO(raw)) as image:
            result, _ = RapidOCR()(np.asarray(image.convert("RGB")))
        return "\n".join(str(item[1]).strip() for item in (result or []) if len(item) >= 3)


def _text_to_blocks(text: str, *, page: int | None = None, source: str) -> list[DocumentBlock]:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not lines:
        return []
    table_like = sum(
        "|" in line or "\t" in line or re.search(r"\s{2,}", line) is not None
        for line in lines
    )
    if len(lines) >= 2 and table_like >= max(2, len(lines) // 2):
        return [
            DocumentBlock(
                "table",
                "\n".join(lines),
                page=page,
                table_id=f"page-{page or 0}-table-1",
                row_start=1,
                row_end=len(lines),
                source=source,
            )
        ]
    return [DocumentBlock("paragraph", "\n".join(lines), page=page, source=source)]
